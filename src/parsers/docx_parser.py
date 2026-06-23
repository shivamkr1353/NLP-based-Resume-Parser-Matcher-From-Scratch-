"""
DOCX Parser — Extract text from Word document resumes
=======================================================
Uses python-docx (allowed utility library) to extract text from .docx files.

Note: This is a UTILITY — not part of the from-scratch ML/NLP pipeline.
"""

import io


def extract_text_from_docx(file_input):
    """
    Extract all text content from a DOCX file.

    Extracts text from paragraphs and tables.

    Args:
        file_input: Either a file path (str) or a file-like object
                    (e.g., from Streamlit's file_uploader)

    Returns:
        str: Extracted text from all paragraphs and tables

    Raises:
        RuntimeError: If extraction fails
    """
    try:
        from docx import Document

        if isinstance(file_input, str):
            doc = Document(file_input)
        elif hasattr(file_input, 'read'):
            content = file_input.read()
            if hasattr(file_input, 'seek'):
                file_input.seek(0)
            doc = Document(io.BytesIO(content))
        else:
            doc = Document(io.BytesIO(file_input))

        text_parts = []

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        return '\n'.join(text_parts)

    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX: {e}")
